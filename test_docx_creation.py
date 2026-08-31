#!/usr/bin/env python
"""Create test DOCX file for validation"""
from docx import Document
import os

# Ensure directory exists
os.makedirs('backend/data/workspaces', exist_ok=True)

# Create a test DOCX file
doc = Document()

# Add title
title = doc.add_heading('Lab Assignment Test Document', 0)

# Add some content
doc.add_heading('Introduction', level=1)
doc.add_paragraph(
    'This is a test document for validating DOCX processing in the Sovereign AI Workbench. '
    'It contains various sections to ensure proper text extraction and chunking.'
)

doc.add_heading('Objective', level=1)
doc.add_paragraph(
    'To verify that uploaded DOCX files are properly processed and indexed in the local RAG system '
    'rather than returning placeholder text.'
)

doc.add_heading('Key Findings', level=1)
findings = doc.add_paragraph()
findings.add_run('Finding 1: ').bold = True
findings.add_run('System should extract actual document content\n')
findings.add_run('Finding 2: ').bold = True
findings.add_run('RAG search should return ingested documents\n')
findings.add_run('Finding 3: ').bold = True
findings.add_run('Reports should reference actual document content')

doc.add_heading('Technical Details', level=1)
doc.add_paragraph('Processing layers:')
doc.add_paragraph('1. Document Processor extracts text from DOCX files using python-docx', style='List Bullet')
doc.add_paragraph('2. Text is chunked into 500-word segments for RAG indexing', style='List Bullet')
doc.add_paragraph('3. Vector Store stores chunks for similarity search', style='List Bullet')
doc.add_paragraph('4. Orchestrator uses RAG results in task execution', style='List Bullet')

# Add a table
table = doc.add_table(rows=3, cols=2)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Component'
hdr_cells[1].text = 'Status'
hdr_cells = table.rows[1].cells
hdr_cells[0].text = 'Document Processor'
hdr_cells[1].text = 'Fixed - Handles DOCX/XLSX'
hdr_cells = table.rows[2].cells
hdr_cells[0].text = 'RAG Vector Store'
hdr_cells[1].text = 'Fixed - Returns real documents'

doc.add_heading('Conclusion', level=1)
doc.add_paragraph(
    'This document serves as a validation artifact to confirm that the system properly '
    'processes and retrieves real document content from uploaded files.'
)

# Save the document
doc.save('backend/data/workspaces/Lab_Assignment_Test.docx')
print("✓ Created test DOCX file: backend/data/workspaces/Lab_Assignment_Test.docx")
